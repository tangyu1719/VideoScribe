#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试 7 类文件（图片/PDF/Word/Markdown/CSV/音频/视频），
并把输入样例与解析产物统一落盘到本地文件夹。
"""

from __future__ import annotations

import csv
import json
import math
import struct
import subprocess
import wave
from datetime import datetime
from pathlib import Path

from document_processor import DocumentProcessor
from ffmpeg_path import ensure_ffmpeg_path


def _write_png(path: Path) -> None:
    from PIL import Image, ImageDraw

    img = Image.new("RGB", (640, 240), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((20, 80), "MULTIMODAL_IMAGE_SENTINEL_001", fill=(0, 0, 0))
    img.save(path)


def _write_pdf(path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(path), pagesize=A4)
    c.setFont("Helvetica", 12)
    c.drawString(72, 780, "MULTIMODAL_PDF_SENTINEL_002")
    c.drawString(72, 760, "This is a PDF content test.")
    c.showPage()
    c.save()


def _write_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Word 测试文档", level=1)
    doc.add_paragraph("MULTIMODAL_DOCX_SENTINEL_003")
    doc.add_paragraph("这是用于多模态解析测试的 Word 正文。")
    doc.save(str(path))


def _write_md(path: Path) -> None:
    path.write_text(
        "# Markdown 测试\n\n"
        "MULTIMODAL_MD_SENTINEL_004\n\n"
        "- 条目A\n"
        "- 条目B\n",
        encoding="utf-8",
    )


def _write_csv(path: Path) -> None:
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["id", "name", "note"])
        writer.writerow([1, "alice", "MULTIMODAL_CSV_SENTINEL_005"])
        writer.writerow([2, "bob", "csv content test"])


def _write_wav(path: Path, duration_sec: float = 2.0, sample_rate: int = 16000) -> None:
    # 生成简短正弦波，确保音频解析链路可跑
    amp = 14000
    freq = 440.0
    total = int(duration_sec * sample_rate)
    with wave.open(str(path), "wb") as wf:
        wf.setnchannels(1)
        wf.setsampwidth(2)
        wf.setframerate(sample_rate)
        for i in range(total):
            v = int(amp * math.sin(2 * math.pi * freq * (i / sample_rate)))
            wf.writeframes(struct.pack("<h", v))


def _write_mp4_with_audio(video_path: Path, wav_path: Path) -> None:
    ensure_ffmpeg_path()
    # 用黑屏 + wav 生成一个最小可处理 mp4
    cmd = [
        "ffmpeg",
        "-y",
        "-f",
        "lavfi",
        "-i",
        "color=c=black:s=640x360:d=2",
        "-i",
        str(wav_path),
        "-shortest",
        "-c:v",
        "libx264",
        "-c:a",
        "aac",
        str(video_path),
    ]
    subprocess.run(cmd, check=True, capture_output=True, text=True)


def main() -> None:
    base = Path(__file__).resolve().parent
    out_root = base / "output" / f"all_file_types_test_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    sample_dir = out_root / "samples"
    result_dir = out_root / "results"
    sample_dir.mkdir(parents=True, exist_ok=True)
    result_dir.mkdir(parents=True, exist_ok=True)

    png = sample_dir / "sample_image.png"
    pdf = sample_dir / "sample_pdf.pdf"
    docx = sample_dir / "sample_docx.docx"
    md = sample_dir / "sample_md.md"
    csv_file = sample_dir / "sample_csv.csv"
    wav_file = sample_dir / "sample_audio.wav"
    mp4_file = sample_dir / "sample_video.mp4"

    _write_png(png)
    _write_pdf(pdf)
    _write_docx(docx)
    _write_md(md)
    _write_csv(csv_file)
    _write_wav(wav_file)
    _write_mp4_with_audio(mp4_file, wav_file)

    processor = DocumentProcessor()
    cases = [
        ("图片", png),
        ("PDF", pdf),
        ("Word", docx),
        ("Markdown", md),
        ("CSV", csv_file),
        ("音频", wav_file),
        ("视频", mp4_file),
    ]

    summary = []
    for name, file_path in cases:
        result = processor.process(str(file_path))
        text = (result.content.text or "") if result.content else ""
        metadata = (result.content.metadata or {}) if result.content else {}
        tables = (result.content.tables or []) if result.content else []

        case_payload = {
            "case": name,
            "input_file": str(file_path),
            "success": bool(result.success),
            "doc_type": result.doc_type.value if result.doc_type else "unknown",
            "error": result.error,
            "processing_time_sec": result.processing_time,
            "text_preview": text[:500],
            "text_length": len(text),
            "metadata": metadata,
            "table_count": len(tables),
        }
        summary.append(case_payload)

        out_json = result_dir / f"{name}.json"
        out_txt = result_dir / f"{name}.txt"
        out_json.write_text(json.dumps(case_payload, ensure_ascii=False, indent=2), encoding="utf-8")
        out_txt.write_text(text, encoding="utf-8")

    summary_file = result_dir / "summary.json"
    summary_file.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")

    ok_count = sum(1 for x in summary if x["success"])
    print(f"输出目录: {out_root}")
    print(f"测试完成: {ok_count}/{len(summary)} 成功")
    for item in summary:
        print(f"- {item['case']}: success={item['success']} text_length={item['text_length']} error={item['error']}")


if __name__ == "__main__":
    main()

