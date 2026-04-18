#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MinerU 文档处理器
基于MinerU技术的PDF/文档解析工具
支持：PDF转Markdown、表格识别、公式识别、OCR
"""

import os
import json
import logging
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
import sys
import shutil
import threading
import tempfile

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def _try_enable_local_mineru_models():
    """
    若用户已通过 `F:\\java\\MinerU` 下载模型并生成 `~/mineru.json`，
    自动把 MINERU_MODEL_SOURCE 设为 local，避免每次联网拉取。
    """
    try:
        cfg = os.path.join(os.path.expanduser("~"), "mineru.json")
        if not os.path.exists(cfg):
            return
        data = json.load(open(cfg, "r", encoding="utf-8"))
        models_dir = (data.get("models-dir") or {}).get("pipeline")
        if models_dir and os.path.exists(models_dir):
            os.environ.setdefault("MINERU_MODEL_SOURCE", "local")
    except Exception:
        pass

_try_enable_local_mineru_models()


def _patch_fast_langdetect_model_path():
    """
    fast_langdetect 在某些 Windows 中文路径下加载 lid.176.ftz 会失败。
    这里把模型复制到 ASCII 路径并重定向 LOCAL_SMALL_MODEL_PATH，避免反复报错和回退。
    """
    try:
        import fast_langdetect
        from fast_langdetect.ft_detect import infer as ft_infer

        src = Path(fast_langdetect.__file__).parent / "ft_detect" / "resources" / "lid.176.ftz"
        if not src.exists():
            return

        dst_dir = Path("C:/mineru-fastlang-cache")
        dst_dir.mkdir(parents=True, exist_ok=True)
        dst = dst_dir / "lid.176.ftz"
        if not dst.exists() or dst.stat().st_size != src.stat().st_size:
            shutil.copy2(src, dst)

        ft_infer.LOCAL_SMALL_MODEL_PATH = dst
        os.environ.setdefault("FTLANG_CACHE", str(dst_dir))

        # 降低 fast_langdetect 警告刷屏，避免拖慢日志渲染
        logging.getLogger("fast_langdetect").setLevel(logging.ERROR)
        logging.getLogger("fast_langdetect.ft_detect.infer").setLevel(logging.ERROR)
    except Exception:
        pass


class MinerUStatus(Enum):
    """处理状态"""
    PENDING = "pending"
    PROCESSING = "processing"
    SUCCESS = "success"
    FAILED = "failed"


@dataclass
class MinerUResult:
    """MinerU处理结果"""
    success: bool
    content: str = ""
    markdown: str = ""
    metadata: Dict = None
    images: List[str] = None
    tables: List[Dict] = None
    error: str = ""
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}
        if self.images is None:
            self.images = []
        if self.tables is None:
            self.tables = []


class MinerUProcessor:
    """
    MinerU文档处理器
    
    功能：
    1. PDF转Markdown（保留格式、表格、公式）
    2. 图片OCR识别
    3. 文档结构化提取
    4. 批量处理支持
    """
    
    # 全局预热状态（避免多个页面/模块重复预热）
    _global_warmed_up = False
    _global_warmup_method = ""

    def __init__(self, output_dir: str = None):
        """
        初始化MinerU处理器
        
        Args:
            output_dir: 输出目录，默认为当前目录下的output文件夹
        """
        self.output_dir = output_dir or os.path.join(os.path.dirname(__file__), "output")
        os.makedirs(self.output_dir, exist_ok=True)
        self._warmup_lock = threading.Lock()
        self._warmed_up = False
        
        # 检查MinerU是否安装
        self.mineru_available = self._check_mineru()
        
        if self.mineru_available:
            logger.info("[MinerU] 处理器初始化成功")
        else:
            logger.warning("[MinerU] 未安装，将使用备用处理方案")

    def warmup(self, force: bool = False) -> MinerUResult:
        """
        预热 MinerU：在服务启动时提前加载模型，减少首次 PDF 处理等待。
        返回 MinerUResult，便于上层界面输出日志。
        """
        with self._warmup_lock:
            if MinerUProcessor._global_warmed_up and not force:
                return MinerUResult(
                    success=True,
                    content="",
                    markdown="",
                    metadata={"method": MinerUProcessor._global_warmup_method or "warmup_cached"},
                    error="",
                )
            if self._warmed_up and not force:
                return MinerUResult(
                    success=True,
                    content="",
                    markdown="",
                    metadata={"method": "warmup_cached"},
                    error="",
                )

            if not self.mineru_available:
                return MinerUResult(success=False, error="MinerU 不可用，无法预热")

            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.pdfgen import canvas

                with tempfile.TemporaryDirectory(prefix="mineru-warmup-") as td:
                    test_pdf = os.path.join(td, "warmup.pdf")
                    c = canvas.Canvas(test_pdf, pagesize=A4)
                    c.setFont("Helvetica", 12)
                    # 预热样本使用更接近真实输入的多行中英混排文本，减少“样本为空”概率
                    c.drawString(72, 790, "MinerU Warmup Document")
                    c.drawString(72, 770, "这是一段用于模型预热的中文文本，包含数字 12345。")
                    c.drawString(72, 750, "Warmup checks parser bootstrap only, not extraction quality.")
                    c.showPage()
                    c.save()

                    result = self.process_pdf(test_pdf, is_warmup=True)

                # 只要预热流程已跑完一次就标记完成，避免重复首启预热拖慢启动。
                # 若需要强制重跑，调用 warmup(force=True)。
                if result.success:
                    self._warmed_up = True
                    MinerUProcessor._global_warmed_up = True
                    method = (result.metadata or {}).get("method", "warmup_cached")
                    MinerUProcessor._global_warmup_method = method
                    if method == "fallback":
                        logger.info(
                            "[MinerU] 预热已完成（warmup 使用备用解析）；正式 PDF 仍会优先尝试 mineru_pipeline"
                        )
                return result
            except Exception as e:
                return MinerUResult(success=False, error=f"预热异常: {e}")
    
    def _check_mineru(self) -> bool:
        """检查MinerU是否已安装"""
        try:
            # 优先：项目投入使用的 mineru（本地源码常见路径 F:\java\MinerU）
            try:
                import mineru  # noqa: F401
                logger.info("[MinerU] mineru 模块已安装")
                return True
            except ImportError:
                pass

            # 兼容：旧实现 magic_pdf
            try:
                import magic_pdf  # noqa: F401
                logger.info("[MinerU] magic_pdf 模块已安装")
                return True
            except ImportError:
                pass

            local_src = r"F:\java\MinerU"
            if os.path.isdir(local_src) and local_src not in sys.path:
                sys.path.insert(0, local_src)
                try:
                    import mineru  # noqa: F401
                    logger.info(f"[MinerU] 已从本地目录加载 mineru: {local_src}")
                    return True
                except ImportError:
                    pass

            logger.warning("[MinerU] 未检测到 mineru/magic_pdf，将使用备用处理方案")
            return False
        except Exception as e:
            logger.warning(f"[MinerU] 检测失败，将使用备用处理方案: {e}")
            return False
    
    def process_pdf(self, pdf_path: str, method: str = "auto", is_warmup: bool = False) -> MinerUResult:
        """
        处理PDF文件
        
        Args:
            pdf_path: PDF文件路径
            method: 处理方法 (auto/ocr/txt)
        
        Returns:
            MinerUResult: 处理结果
        """
        if not os.path.exists(pdf_path):
            return MinerUResult(success=False, error=f"文件不存在: {pdf_path}")
        
        if not self.mineru_available:
            return self._process_pdf_fallback(pdf_path)
        
        try:
            try:
                _patch_fast_langdetect_model_path()
                # 你本地“投入使用”的 MinerU 是 F:\java\MinerU 的 `mineru` 包，不是 pip 的 magic_pdf。
                # 这里直接调用 mineru 的 pipeline 链路产出 md，并读取最终 md 内容。
                if r"F:\java\MinerU" not in sys.path and os.path.isdir(r"F:\java\MinerU"):
                    sys.path.insert(0, r"F:\java\MinerU")

                from mineru.cli.common import (  # type: ignore
                    read_fn,
                    prepare_env,
                    convert_pdf_bytes_to_bytes_by_pypdfium2,
                )
                from mineru.data.data_reader_writer import FileBasedDataWriter  # type: ignore
                from mineru.backend.pipeline.pipeline_analyze import doc_analyze as pipeline_doc_analyze  # type: ignore
                from mineru.backend.pipeline.model_json_to_middle_json import (  # type: ignore
                    result_to_middle_json as pipeline_result_to_middle_json,
                )
                from mineru.backend.pipeline.pipeline_middle_json_mkcontent import union_make as pipeline_union_make  # type: ignore
                from mineru.utils.enum_class import MakeMode  # type: ignore

                logger.info(f"[MinerU] 使用本地 mineru(pipeline) 解析 PDF: {pdf_path}")

                pdf_bytes = read_fn(Path(pdf_path))
                pdf_file_name = Path(pdf_path).stem
                local_image_dir, local_md_dir = prepare_env(self.output_dir, pdf_file_name, "auto")
                image_writer = FileBasedDataWriter(local_image_dir)
                md_writer = FileBasedDataWriter(local_md_dir)

                pdf_bytes_processed = convert_pdf_bytes_to_bytes_by_pypdfium2(pdf_bytes, 0, None)
                # 性能开关：默认关闭公式/表格识别以加速普通文本型 PDF
                formula_enable = os.environ.get("MINERU_PDF_FORMULA_ENABLE", "false").lower() == "true"
                table_enable = os.environ.get("MINERU_PDF_TABLE_ENABLE", "false").lower() == "true"
                infer_results, all_image_lists, all_pdf_docs, lang_list, ocr_enabled_list = pipeline_doc_analyze(
                    [pdf_bytes_processed],
                    ["ch"],
                    parse_method="auto",
                    formula_enable=formula_enable,
                    table_enable=table_enable,
                )

                markdown_content = ""
                for idx, model_list in enumerate(infer_results):
                    import copy

                    model_json = copy.deepcopy(model_list)
                    images_list = all_image_lists[idx]
                    pdf_doc = all_pdf_docs[idx]
                    _lang = lang_list[idx]
                    _ocr_enable = ocr_enabled_list[idx]

                    middle_json = pipeline_result_to_middle_json(
                        model_list,
                        images_list,
                        pdf_doc,
                        image_writer,
                        _lang,
                        _ocr_enable,
                        True,
                    )
                    pdf_info = middle_json["pdf_info"]
                    markdown_content = pipeline_union_make(
                        pdf_info, MakeMode.MM_MD, str(os.path.basename(local_image_dir))
                    )
                    md_writer.write_string(f"{pdf_file_name}.md", markdown_content)
                    md_writer.write_string(f"{pdf_file_name}_middle.json", json.dumps(middle_json, ensure_ascii=False))
                    md_writer.write_string(f"{pdf_file_name}_model.json", json.dumps(model_json, ensure_ascii=False))

                markdown_content = (markdown_content or "").strip()
                if markdown_content:
                    return MinerUResult(
                        success=True,
                        content=markdown_content,
                        markdown=markdown_content,
                        metadata={
                            "file_name": pdf_file_name,
                            "file_path": pdf_path,
                            "output_path": local_md_dir,
                            "method": "mineru_pipeline",
                        },
                    )
                raise RuntimeError("mineru 输出为空")
            except Exception as e:
                # 预热样本（warmup.pdf）为空是常见情况，不应误导为“功能失败”。
                # 真实业务文件仍保持 warning 级别，便于排障。
                if is_warmup and ("输出为空" in str(e) or "empty" in str(e).lower()):
                    logger.info(
                        "[MinerU] 预热样本输出为空（仅预热样本），已回退备用解析完成预热；"
                        "正式处理仍会先走 mineru_pipeline"
                    )
                else:
                    logger.warning(f"[MinerU] 本地 mineru 解析失败，将回退备用方案（不阻塞投入使用）: {e}")
                return self._process_pdf_fallback(pdf_path)

            # （保留）若未来环境装了 magic_pdf，可在这里补充回退分支；
            # 当前生产环境以 mineru + 备用解析为主，避免因未安装 magic_pdf 而失败。
            
            logger.info(f"[MinerU] 开始处理PDF: {pdf_path}")
            
            # 生成输出目录
            file_name = Path(pdf_path).stem
            output_path = os.path.join(self.output_dir, file_name)
            os.makedirs(output_path, exist_ok=True)
            
            # 读取PDF
            pdf_bytes = open(pdf_path, "rb").read()
            
            # 选择解析方法
            if method == "ocr":
                parse_method = SupportedPdfParseMethod.OCR
            elif method == "txt":
                parse_method = SupportedPdfParseMethod.TXT
            else:
                parse_method = SupportedPdfParseMethod.AUTO
            
            # 创建数据集
            dataset = PymuDocDataset(pdf_bytes)
            
            # 解析文档
            parsed_result = dataset.apply(doc_analyze, ocr=True)
            
            # 提取内容
            markdown_content = parsed_result.get_markdown()
            
            # 保存Markdown
            md_path = os.path.join(output_path, f"{file_name}.md")
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            # 提取图片
            images = []
            for img_info in parsed_result.get_images():
                img_path = os.path.join(output_path, "images", img_info["filename"])
                images.append(img_path)
            
            # 提取表格
            tables = parsed_result.get_tables()
            
            logger.info(f"[MinerU] PDF处理完成: {file_name}")
            
            return MinerUResult(
                success=True,
                content=markdown_content,
                markdown=markdown_content,
                metadata={
                    "file_name": file_name,
                    "file_path": pdf_path,
                    "output_path": output_path,
                    "page_count": len(parsed_result),
                    "image_count": len(images),
                    "table_count": len(tables)
                },
                images=images,
                tables=tables
            )
            
        except Exception as e:
            logger.error(f"[MinerU] PDF处理失败: {e}")
            return MinerUResult(success=False, error=str(e))
    
    def _process_pdf_fallback(self, pdf_path: str) -> MinerUResult:
        """PDF处理备用方案（使用PyPDF2）"""
        try:
            import PyPDF2
            
            logger.info(f"[MinerU] 使用备用方案处理PDF: {pdf_path}")
            
            file_name = Path(pdf_path).stem
            text_content = []
            
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(reader.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_content.append(f"## 第{page_num}页\n\n{text}")
            
            markdown_content = "\n\n".join(text_content)
            
            # 保存Markdown
            output_path = os.path.join(self.output_dir, file_name)
            os.makedirs(output_path, exist_ok=True)
            md_path = os.path.join(output_path, f"{file_name}.md")
            
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            return MinerUResult(
                success=True,
                content=markdown_content,
                markdown=markdown_content,
                metadata={
                    "file_name": file_name,
                    "file_path": pdf_path,
                    "output_path": output_path,
                    "page_count": len(reader.pages),
                    "method": "fallback"
                }
            )
            
        except Exception as e:
            logger.error(f"[MinerU] 备用处理失败: {e}")
            return MinerUResult(success=False, error=str(e))
    
    def process_image(self, image_path: str, ocr: bool = True) -> MinerUResult:
        """
        处理图片（OCR识别）
        
        Args:
            image_path: 图片路径
            ocr: 是否启用OCR
        
        Returns:
            MinerUResult: 处理结果
        """
        if not os.path.exists(image_path):
            return MinerUResult(success=False, error=f"文件不存在: {image_path}")
        
        try:
            logger.info(f"[MinerU] 处理图片: {image_path}")
            
            # 尝试使用PaddleOCR
            try:
                from paddleocr import PaddleOCR
                ocr_engine = PaddleOCR(use_angle_cls=True, lang='ch')
                result = ocr_engine.ocr(image_path, cls=True)
                
                text_lines = []
                for line in result[0]:
                    if line:
                        text_lines.append(line[1][0])
                
                content = "\n".join(text_lines)
                
                return MinerUResult(
                    success=True,
                    content=content,
                    markdown=f"![图片]({image_path})\n\n## OCR识别结果\n\n{content}",
                    metadata={
                        "file_name": Path(image_path).name,
                        "file_path": image_path,
                        "ocr_enabled": True,
                        "text_lines": len(text_lines)
                    }
                )
                
            except ImportError:
                logger.warning("[MinerU] PaddleOCR未安装，跳过OCR")
                return MinerUResult(
                    success=True,
                    content="",
                    markdown=f"![图片]({image_path})",
                    metadata={
                        "file_name": Path(image_path).name,
                        "file_path": image_path,
                        "ocr_enabled": False
                    }
                )
                
        except Exception as e:
            logger.error(f"[MinerU] 图片处理失败: {e}")
            return MinerUResult(success=False, error=str(e))
    
    def process_document(self, file_path: str) -> MinerUResult:
        """
        处理文档（自动识别类型）
        
        Args:
            file_path: 文档路径
        
        Returns:
            MinerUResult: 处理结果
        """
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return self.process_pdf(file_path)
        elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp']:
            return self.process_image(file_path)
        elif ext in ['.docx', '.doc']:
            return self._process_word(file_path)
        elif ext in ['.md', '.markdown']:
            return self._process_markdown(file_path)
        else:
            return MinerUResult(success=False, error=f"不支持的文件类型: {ext}")
    
    def _process_word(self, docx_path: str) -> MinerUResult:
        """处理Word文档"""
        try:
            from docx import Document
            
            logger.info(f"[MinerU] 处理Word文档: {docx_path}")
            
            doc = Document(docx_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            content = "\n\n".join(paragraphs)
            markdown_content = f"# {Path(docx_path).stem}\n\n{content}"
            
            return MinerUResult(
                success=True,
                content=content,
                markdown=markdown_content,
                metadata={
                    "file_name": Path(docx_path).name,
                    "file_path": docx_path,
                    "paragraph_count": len(paragraphs)
                }
            )
            
        except ImportError:
            return MinerUResult(success=False, error="python-docx未安装")
        except Exception as e:
            return MinerUResult(success=False, error=str(e))
    
    def _process_markdown(self, md_path: str) -> MinerUResult:
        """处理Markdown文件"""
        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return MinerUResult(
                success=True,
                content=content,
                markdown=content,
                metadata={
                    "file_name": Path(md_path).name,
                    "file_path": md_path,
                    "char_count": len(content)
                }
            )
        except Exception as e:
            return MinerUResult(success=False, error=str(e))
    
    def batch_process(self, file_paths: List[str]) -> List[MinerUResult]:
        """批量处理文件"""
        results = []
        for file_path in file_paths:
            result = self.process_document(file_path)
            results.append(result)
        return results
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式"""
        return ['.pdf', '.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp', 
                '.docx', '.doc', '.md', '.markdown']


# 便捷函数
def process_with_mineru(file_path: str, output_dir: str = None) -> MinerUResult:
    """便捷函数：使用MinerU处理文件"""
    processor = MinerUProcessor(output_dir=output_dir)
    return processor.process_document(file_path)


if __name__ == "__main__":
    # 测试
    print("=" * 60)
    print("MinerU 文档处理器测试")
    print("=" * 60)
    
    processor = MinerUProcessor()
    print(f"MinerU可用: {processor.mineru_available}")
    print(f"支持格式: {processor.get_supported_formats()}")
