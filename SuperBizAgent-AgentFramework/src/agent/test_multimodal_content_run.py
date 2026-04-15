#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
多模态文件内容跑通测试（必须真实读取到正文）

覆盖：
- Markdown (.md)
- Word (.docx)
- PDF (.pdf)

目标：
1) MinerUProcessor.process_document 能成功处理三类文件
2) 输出内容包含我们写入的“哨兵文本”，证明不是空跑
"""

from __future__ import annotations

import os
import tempfile
from pathlib import Path


def _write_md(p: Path) -> None:
    p.write_text(
        "# 多模态测试\n\n"
        "哨兵文本: MULTIMODAL_MD_SENTINEL_123\n\n"
        "- 列表项A\n"
        "- 列表项B\n",
        encoding="utf-8",
    )


def _write_docx(p: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("多模态测试", level=1)
    doc.add_paragraph("哨兵文本: MULTIMODAL_DOCX_SENTINEL_456")
    doc.add_paragraph("第二段：用于验证段落拼接。")
    doc.save(str(p))


def _write_pdf(p: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas

    # 使用 CID 字体保证中文可绘制；PyPDF2 的 extract_text 对“画出来的字”提取能力
    # 依赖 PDF 内部编码。这里用英文哨兵确保可提取，中文用于人工肉眼校验。
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    c = canvas.Canvas(str(p), pagesize=A4)
    c.setFont("STSong-Light", 14)
    c.drawString(72, 800, "多模态测试 PDF（中文行）")
    c.setFont("Helvetica", 12)
    c.drawString(72, 780, "Sentinel: MULTIMODAL_PDF_SENTINEL_789")
    c.drawString(72, 760, "Second line: hello pdf")
    c.showPage()
    c.save()


def _assert_contains(haystack: str, needle: str, label: str) -> None:
    if needle not in haystack:
        raise AssertionError(f"{label} 未包含哨兵文本: {needle!r}，实际内容前200字符: {haystack[:200]!r}")

def _preview_ascii(s: str, limit: int = 160) -> str:
    """
    Windows 控制台可能是 GBK，直接 print 中文会报 UnicodeEncodeError。
    这里把字符串转成 ascii 安全的转义形式，保证日志可输出。
    """
    return (s[:limit]).encode("unicode_escape", errors="backslashreplace").decode("ascii", errors="ignore")


def main() -> None:
    here = Path(__file__).resolve().parent
    os.chdir(str(here))

    from mineru_processor import MinerUProcessor

    proc = MinerUProcessor(output_dir=str(here / "output"))

    with tempfile.TemporaryDirectory(prefix="multimodal-content-") as td:
        td = Path(td)
        md_path = td / "case.md"
        docx_path = td / "case.docx"
        pdf_path = td / "case.pdf"

        _write_md(md_path)
        _write_docx(docx_path)
        _write_pdf(pdf_path)

        cases = [
            ("md", md_path, "MULTIMODAL_MD_SENTINEL_123"),
            ("docx", docx_path, "MULTIMODAL_DOCX_SENTINEL_456"),
            ("pdf", pdf_path, "MULTIMODAL_PDF_SENTINEL_789"),
        ]

        print("=" * 70)
        print("多模态文件内容跑通测试（MinerUProcessor.process_document）")
        print(f"MinerU 可用: {proc.mineru_available}")
        print("=" * 70)

        for kind, path, sentinel in cases:
            print(f"\n[{kind.upper()}] 处理文件: {path}")
            res = proc.process_document(str(path))
            print(f"  success={res.success}")
            if not res.success:
                raise RuntimeError(f"{kind} 处理失败: {res.error}")

            text = (res.content or res.markdown or "").strip()
            print(f"  extracted_chars={len(text)}")
            print(f"  extracted_preview_ascii={_preview_ascii(text)}")
            _assert_contains(text, sentinel, kind)

        print("\nALL_OK")


if __name__ == "__main__":
    main()

