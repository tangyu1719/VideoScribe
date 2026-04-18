#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
多模态文档处理模块
支持文件类型：图片、PDF、DOCX、MD、CSV、音频、视频、网页链接
"""

import os
import io
import re
import json
import hashlib
import tempfile
from pathlib import Path
from typing import Optional, Dict, Any, List, Union, BinaryIO
from dataclasses import dataclass, field
from enum import Enum
import subprocess

from ffmpeg_path import ensure_ffmpeg_path

# 文件类型枚举
class DocumentType(Enum):
    # 链接类型
    WEB_LINK = "web_link"           # 网页链接
    XIAOHONGSHU = "xiaohongshu"     # 小红书
    DOUYIN = "douyin"               # 抖音
    BILIBILI = "bilibili"           # B站
    
    # 文件类型
    IMAGE = "image"                 # 图片 (jpg, png, gif, webp)
    PDF = "pdf"                     # PDF文档
    DOCX = "docx"                   # Word文档
    MD = "markdown"                 # Markdown文档
    CSV = "csv"                     # CSV表格
    AUDIO = "audio"                 # 音频 (mp3, wav, m4a)
    VIDEO = "video"                 # 视频 (mp4, avi, mov)
    UNKNOWN = "unknown"             # 未知类型

# 文件扩展名映射
EXTENSION_MAP = {
    # 图片
    '.jpg': DocumentType.IMAGE, '.jpeg': DocumentType.IMAGE,
    '.png': DocumentType.IMAGE, '.gif': DocumentType.IMAGE,
    '.webp': DocumentType.IMAGE, '.bmp': DocumentType.IMAGE,
    
    # 文档
    '.pdf': DocumentType.PDF,
    '.docx': DocumentType.DOCX, '.doc': DocumentType.DOCX,
    '.md': DocumentType.MD, '.markdown': DocumentType.MD,
    '.csv': DocumentType.CSV,
    
    # 音频
    '.mp3': DocumentType.AUDIO, '.wav': DocumentType.AUDIO,
    '.m4a': DocumentType.AUDIO, '.flac': DocumentType.AUDIO,
    '.ogg': DocumentType.AUDIO, '.aac': DocumentType.AUDIO,
    
    # 视频
    '.mp4': DocumentType.VIDEO, '.avi': DocumentType.VIDEO,
    '.mov': DocumentType.VIDEO, '.mkv': DocumentType.VIDEO,
    '.flv': DocumentType.VIDEO, '.wmv': DocumentType.VIDEO,
}

@dataclass
class ExtractedContent:
    """提取的内容结构"""
    text: str = ""                           # 文本内容
    images: List[Dict[str, Any]] = field(default_factory=list)  # 图片列表
    tables: List[List[List[str]]] = field(default_factory=list)  # 表格数据
    metadata: Dict[str, Any] = field(default_factory=dict)       # 元数据
    chunks: List[str] = field(default_factory=list)              # 文本分块

@dataclass
class ProcessingResult:
    """处理结果"""
    success: bool = False
    doc_type: DocumentType = DocumentType.UNKNOWN
    content: ExtractedContent = field(default_factory=ExtractedContent)
    error: Optional[str] = None
    file_path: Optional[str] = None
    file_size: int = 0
    processing_time: float = 0.0

class BaseParser:
    """基础解析器"""
    
    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
    
    def parse(self, file_path: str, **kwargs) -> ProcessingResult:
        """解析文件"""
        raise NotImplementedError

class ImageParser(BaseParser):
    """图片解析器 - OCR提取文字"""
    
    def parse(self, file_path: str, **kwargs) -> ProcessingResult:
        import time
        start_time = time.time()
        
        try:
            # 使用PIL读取图片
            from PIL import Image
            
            img = Image.open(file_path)
            
            # 尝试OCR提取文字
            text = ""
            try:
                import pytesseract
                # 设置中文识别
                text = pytesseract.image_to_string(img, lang='chi_sim+eng')
            except Exception as e:
                print(f"[ImageParser] OCR失败: {e}")
            
            # 构建结果
            result = ProcessingResult(
                success=True,
                doc_type=DocumentType.IMAGE,
                content=ExtractedContent(
                    text=text.strip(),
                    images=[{
                        "path": file_path,
                        "size": img.size,
                        "mode": img.mode
                    }],
                    metadata={
                        "format": img.format,
                        "size": img.size,
                        "mode": img.mode
                    }
                ),
                file_path=file_path,
                file_size=os.path.getsize(file_path),
                processing_time=time.time() - start_time
            )
            
            return result
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                doc_type=DocumentType.IMAGE,
                error=str(e),
                file_path=file_path
            )

class PDFParser(BaseParser):
    """PDF解析器"""
    
    def parse(self, file_path: str, **kwargs) -> ProcessingResult:
        import time
        start_time = time.time()
        
        try:
            # 尝试使用PyMuPDF (fitz)
            try:
                import fitz  # PyMuPDF
                return self._parse_with_fitz(file_path, start_time)
            except ImportError:
                pass
            
            # 回退到pdfplumber
            try:
                import pdfplumber
                return self._parse_with_pdfplumber(file_path, start_time)
            except ImportError:
                pass
            
            # 最后回退到PyPDF2
            return self._parse_with_pypdf2(file_path, start_time)
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                doc_type=DocumentType.PDF,
                error=str(e),
                file_path=file_path
            )
    
    def _parse_with_fitz(self, file_path: str, start_time: float) -> ProcessingResult:
        """使用PyMuPDF解析"""
        import fitz
        import time
        
        doc = fitz.open(file_path)
        text_parts = []
        images = []
        page_count = len(doc)
        
        for page_num in range(page_count):
            page = doc[page_num]
            
            # 提取文本
            text_parts.append(page.get_text())
            
            # 提取图片
            for img_index, img in enumerate(page.get_images()):
                xref = img[0]
                base_image = doc.extract_image(xref)
                images.append({
                    "page": page_num,
                    "index": img_index,
                    "size": len(base_image["image"])
                })
        
        doc.close()
        
        return ProcessingResult(
            success=True,
            doc_type=DocumentType.PDF,
            content=ExtractedContent(
                text="\n".join(text_parts),
                images=images,
                metadata={
                    "pages": page_count,
                    "format": "PDF"
                }
            ),
            file_path=file_path,
            file_size=os.path.getsize(file_path),
            processing_time=time.time() - start_time
        )
    
    def _parse_with_pdfplumber(self, file_path: str, start_time: float) -> ProcessingResult:
        """使用pdfplumber解析"""
        import pdfplumber
        import time
        
        text_parts = []
        tables = []
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
                
                # 提取表格
                page_tables = page.extract_tables()
                if page_tables:
                    tables.extend(page_tables)
        
        return ProcessingResult(
            success=True,
            doc_type=DocumentType.PDF,
            content=ExtractedContent(
                text="\n".join(text_parts),
                tables=tables,
                metadata={
                    "pages": len(pdf.pages),
                    "format": "PDF"
                }
            ),
            file_path=file_path,
            file_size=os.path.getsize(file_path),
            processing_time=time.time() - start_time
        )
    
    def _parse_with_pypdf2(self, file_path: str, start_time: float) -> ProcessingResult:
        """使用PyPDF2解析"""
        from PyPDF2 import PdfReader
        import time
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            text_parts.append(page.extract_text() or "")
        
        return ProcessingResult(
            success=True,
            doc_type=DocumentType.PDF,
            content=ExtractedContent(
                text="\n".join(text_parts),
                metadata={
                    "pages": len(reader.pages),
                    "format": "PDF"
                }
            ),
            file_path=file_path,
            file_size=os.path.getsize(file_path),
            processing_time=time.time() - start_time
        )

class DOCXParser(BaseParser):
    """Word文档解析器"""
    
    def parse(self, file_path: str, **kwargs) -> ProcessingResult:
        import time
        start_time = time.time()
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # 提取段落
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # 提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            # 提取元数据
            metadata = {
                "paragraphs": len(paragraphs),
                "tables": len(tables)
            }
            
            return ProcessingResult(
                success=True,
                doc_type=DocumentType.DOCX,
                content=ExtractedContent(
                    text="\n".join(paragraphs),
                    tables=tables,
                    metadata=metadata
                ),
                file_path=file_path,
                file_size=os.path.getsize(file_path),
                processing_time=time.time() - start_time
            )
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                doc_type=DocumentType.DOCX,
                error=str(e),
                file_path=file_path
            )

class MarkdownParser(BaseParser):
    """Markdown解析器"""
    
    def parse(self, file_path: str, **kwargs) -> ProcessingResult:
        import time
        start_time = time.time()
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 简单的Markdown处理
            # 移除Markdown标记，提取纯文本
            text = re.sub(r'[#*_`~\[\]\(\)!]', '', content)
            
            return ProcessingResult(
                success=True,
                doc_type=DocumentType.MD,
                content=ExtractedContent(
                    text=text,
                    metadata={
                        "format": "Markdown",
                        "original": content
                    }
                ),
                file_path=file_path,
                file_size=os.path.getsize(file_path),
                processing_time=time.time() - start_time
            )
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                doc_type=DocumentType.MD,
                error=str(e),
                file_path=file_path
            )

class CSVParser(BaseParser):
    """CSV解析器"""
    
    def parse(self, file_path: str, **kwargs) -> ProcessingResult:
        import time
        import csv
        start_time = time.time()
        
        try:
            tables = []
            text_parts = []
            
            with open(file_path, 'r', encoding='utf-8-sig') as f:
                reader = csv.reader(f)
                headers = next(reader, None)
                
                if headers:
                    table_data = [headers]
                    for row in reader:
                        table_data.append(row)
                        text_parts.append(", ".join(row))
                    tables.append(table_data)
            
            return ProcessingResult(
                success=True,
                doc_type=DocumentType.CSV,
                content=ExtractedContent(
                    text="\n".join(text_parts),
                    tables=tables,
                    metadata={
                        "format": "CSV",
                        "rows": len(tables[0]) if tables else 0,
                        "columns": len(tables[0][0]) if tables and tables[0] else 0
                    }
                ),
                file_path=file_path,
                file_size=os.path.getsize(file_path),
                processing_time=time.time() - start_time
            )
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                doc_type=DocumentType.CSV,
                error=str(e),
                file_path=file_path
            )

class AudioParser(BaseParser):
    """音频解析器 - 语音转文字"""
    
    def parse(self, file_path: str, **kwargs) -> ProcessingResult:
        import time
        start_time = time.time()
        
        try:
            # 使用whisper进行语音转文字
            from video_downloader import speech_to_text
            
            log_callback = kwargs.get('log_callback')
            progress_callback = kwargs.get('progress_callback')
            llm_config = kwargs.get('llm_config')
            
            transcript_result = speech_to_text(
                file_path,
                log_callback=log_callback,
                progress_callback=progress_callback
            )
            
            if transcript_result:
                return ProcessingResult(
                    success=True,
                    doc_type=DocumentType.AUDIO,
                    content=ExtractedContent(
                        text=transcript_result.get('full_text', ''),
                        metadata={
                            "format": "Audio",
                            "segments": transcript_result.get('segments', []),
                            "ai_summary": transcript_result.get('ai_summary', '')
                        }
                    ),
                    file_path=file_path,
                    file_size=os.path.getsize(file_path),
                    processing_time=time.time() - start_time
                )
            else:
                return ProcessingResult(
                    success=False,
                    doc_type=DocumentType.AUDIO,
                    error="语音转文字失败",
                    file_path=file_path
                )
                
        except Exception as e:
            return ProcessingResult(
                success=False,
                doc_type=DocumentType.AUDIO,
                error=str(e),
                file_path=file_path
            )

class VideoParser(BaseParser):
    """视频解析器 - 提取音频并转文字"""
    
    def parse(self, file_path: str, **kwargs) -> ProcessingResult:
        import time
        start_time = time.time()
        
        try:
            # 首先提取音频
            audio_path = self._extract_audio(file_path)
            
            if not audio_path:
                return ProcessingResult(
                    success=False,
                    doc_type=DocumentType.VIDEO,
                    error="音频提取失败",
                    file_path=file_path
                )
            
            # 使用音频解析器处理
            audio_parser = AudioParser(self.config)
            result = audio_parser.parse(audio_path, **kwargs)
            
            # 清理临时音频文件
            if os.path.exists(audio_path) and audio_path != file_path:
                os.remove(audio_path)
            
            # 更新文档类型为视频
            result.doc_type = DocumentType.VIDEO
            result.file_path = file_path
            result.file_size = os.path.getsize(file_path)
            
            return result
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                doc_type=DocumentType.VIDEO,
                error=str(e),
                file_path=file_path
            )
    
    def _extract_audio(self, video_path: str) -> Optional[str]:
        """从视频中提取音频"""
        try:
            ensure_ffmpeg_path()
            # 创建临时音频文件
            temp_dir = tempfile.gettempdir()
            audio_path = os.path.join(temp_dir, f"{hashlib.md5(video_path.encode()).hexdigest()}.mp3")
            
            # 使用ffmpeg提取音频
            cmd = [
                "ffmpeg",
                "-i", video_path,
                "-vn",  # 不处理视频
                "-acodec", "libmp3lame",
                "-q:a", "2",
                "-y",  # 覆盖输出文件
                audio_path
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            
            if result.returncode == 0 and os.path.exists(audio_path):
                return audio_path
            else:
                print(f"[VideoParser] ffmpeg错误: {result.stderr}")
                return None
                
        except Exception as e:
            print(f"[VideoParser] 提取音频失败: {e}")
            return None

class DocumentProcessor:
    """文档处理器主类"""
    
    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
        self.parsers: Dict[DocumentType, BaseParser] = {
            DocumentType.IMAGE: ImageParser(config),
            DocumentType.PDF: PDFParser(config),
            DocumentType.DOCX: DOCXParser(config),
            DocumentType.MD: MarkdownParser(config),
            DocumentType.CSV: CSVParser(config),
            DocumentType.AUDIO: AudioParser(config),
            DocumentType.VIDEO: VideoParser(config),
        }
    
    def detect_type(self, file_path: str) -> DocumentType:
        """检测文件类型"""
        ext = Path(file_path).suffix.lower()
        return EXTENSION_MAP.get(ext, DocumentType.UNKNOWN)
    
    def process(self, file_path: str, **kwargs) -> ProcessingResult:
        """处理文档"""
        # 检测类型
        doc_type = self.detect_type(file_path)
        
        if doc_type == DocumentType.UNKNOWN:
            return ProcessingResult(
                success=False,
                doc_type=doc_type,
                error=f"不支持的文件类型: {Path(file_path).suffix}",
                file_path=file_path
            )
        
        # 获取解析器
        parser = self.parsers.get(doc_type)
        if not parser:
            return ProcessingResult(
                success=False,
                doc_type=doc_type,
                error=f"未找到对应的解析器: {doc_type.value}",
                file_path=file_path
            )
        
        # 解析文档
        return parser.parse(file_path, **kwargs)
    
    def process_link(self, url: str, **kwargs) -> ProcessingResult:
        """处理链接"""
        # 使用link_analyzer处理链接
        try:
            from link_analyzer import LinkAnalyzer
            
            analyzer = LinkAnalyzer()
            
            # 判断链接类型
            if 'xiaohongshu.com' in url:
                result = analyzer._analyze_xiaohongshu(url)
                doc_type = DocumentType.XIAOHONGSHU
            elif 'douyin.com' in url or 'tiktok.com' in url:
                result = analyzer._analyze_douyin_image(url)
                doc_type = DocumentType.DOUYIN
            elif 'bilibili.com' in url:
                result = analyzer._analyze_bilibili(url)
                doc_type = DocumentType.BILIBILI
            else:
                result = analyzer._analyze_general(url)
                doc_type = DocumentType.WEB_LINK
            
            return ProcessingResult(
                success=True,
                doc_type=doc_type,
                content=ExtractedContent(
                    text=result.get('text_content', ''),
                    images=[{"url": img} for img in result.get('image_links', [])],
                    metadata={
                        "title": result.get('title', ''),
                        "platform": doc_type.value
                    }
                ),
                file_path=url
            )
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                doc_type=DocumentType.WEB_LINK,
                error=str(e),
                file_path=url
            )

# 导出
__all__ = [
    'DocumentProcessor', 'DocumentType', 'ProcessingResult', 'ExtractedContent',
    'ImageParser', 'PDFParser', 'DOCXParser', 'MarkdownParser', 
    'CSVParser', 'AudioParser', 'VideoParser'
]
